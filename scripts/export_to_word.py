#!/usr/bin/env python3
"""
将 Markdown 报告转换为 Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_table(doc, data, headers=None):
    """添加表格"""
    if headers:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
    else:
        table = doc.add_table(rows=len(data), cols=len(data[0]) if data else 0)
        table.style = 'Table Grid'
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    return table

def generate_word_report(markdown_path, output_path):
    """生成 Word 报告"""
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    
    # 标题
    title = doc.add_heading('王茂清 9 周调理总结 + 体检整合报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    # 基本信息
    doc.add_paragraph('报告周期：2025.12.30 - 2026.3.4')
    doc.add_paragraph('体检日期：2025.12.23')
    doc.add_paragraph('体检医院：广东省人民医院')
    doc.add_paragraph()
    
    # 一、主要诊断
    add_heading(doc, '一、主要诊断', level=1)
    diagnoses = [
        '冠状动脉粥样硬化性心脏病',
        '慢性肾脏病 5 期 (CKD5)',
        '2 型糖尿病伴多个并发症',
        '高血压 2 级（很高危）',
        '高钾血症',
        '缺铁性贫血',
        '混合性白内障 + 人工晶体植入'
    ]
    for diag in diagnoses:
        p = doc.add_paragraph(diag, style='List Number')
    
    # 二、体检关键指标
    add_heading(doc, '二、体检关键指标（2025.12.23）', level=1)
    exam_headers = ['项目', '结果', '状态']
    exam_data = [
        ['尿蛋白', '3.0g/L (3+)', '⚠️ 异常'],
        ['总蛋白', '62.4 g/L (偏低)', '⚠️ 异常'],
        ['白蛋白', '37.4 g/L (偏低)', '⚠️ 异常'],
        ['ALT', '56 U/L (偏高)', '⚠️ 异常'],
        ['AST', '44 U/L (偏高)', '⚠️ 异常'],
        ['CHE', '4069 U/L (偏低)', '⚠️ 异常'],
        ['胱抑素 C', '4.63 mg/L (显著升高)', '⚠️ 异常'],
        ['eGFR', '29 ml/min (肾衰竭期)', '⚠️ 异常'],
        ['血红蛋白', '92 g/L (贫血)', '⚠️ 异常'],
        ['空腹血糖', '5.85 mmol/L', '✅ 正常'],
        ['血压', '163/77 mmHg', '⚠️ 异常']
    ]
    add_table(doc, exam_data, exam_headers)
    
    # 三、9 周调理成果对比
    add_heading(doc, '三、9 周调理成果对比', level=1)
    comparison_headers = ['指标', '干预前', '第 9 周', '变化', '状态']
    comparison_data = [
        ['体重', '62.3kg', '56.1kg', '-6.2kg', '✅ 改善'],
        ['空腹血糖', '5.85', '5.6', '-0.25', '✅ 稳定'],
        ['血压', '163/77', '127/65', '-36/-12', '✅ 改善'],
        ['水肿', '双下肢中度水肿', '完全消退', '痊愈', '✅ 改善'],
        ['精神状态', '走路吃力', '状态好很多', '显著提升', '✅ 改善']
    ]
    add_table(doc, comparison_data, comparison_headers)
    
    # 9 周体重变化
    doc.add_paragraph()
    doc.add_paragraph('9 周体重变化曲线：').bold = True
    weight_headers = ['周次', '体重 (kg)', '周变化', '累计变化']
    weight_data = [
        ['第 1 周', '61.45', '-0.85', '-0.85'],
        ['第 2 周', '60.0', '-1.45', '-2.30'],
        ['第 3 周', '58.5', '-1.50', '-3.80'],
        ['第 4 周', '57.5', '-1.00', '-4.80'],
        ['第 5 周', '57.5', '±0.00', '-4.80'],
        ['第 6 周', '58.1', '+0.60', '-4.20'],
        ['第 7 周', '57.5', '-0.60', '-4.80'],
        ['第 8 周', '56.45', '-1.05', '-5.85'],
        ['第 9 周', '56.1', '-0.35', '-6.20']
    ]
    add_table(doc, weight_data, weight_headers)
    
    # 四、9 周调理亮点与进步
    add_heading(doc, '四、9 周调理亮点与进步', level=1)
    highlights = [
        '体重成功减轻 6.2kg，减轻心肾负担',
        '水肿完全消退，从双下肢中度水肿到完全正常',
        '血压显著改善，从 163/77 降至 127/65',
        '精神状态大幅提升，从走路吃力到状态好很多，更加有力气',
        '尿泡沫明显减少，肾功能指标稳定',
        '低血糖问题得到有效控制（换德古胰岛素后）',
        '连续 9 周完整监测，数据记录规范'
    ]
    for highlight in highlights:
        doc.add_paragraph(f'✅ {highlight}', style='List Number')
    
    # 五、需持续关注的问题
    add_heading(doc, '五、需持续关注的问题', level=1)
    concerns = [
        '餐后血糖仍有波动，需加强饮食控制（第 9 周 4/7 超标）',
        '收缩压时有超标，需持续低盐饮食',
        '慢性肾脏病 5 期，需定期复查肾功能（eGFR 29ml/min）',
        '贫血问题，需继续补铁治疗（血红蛋白 92g/L）',
        '多种慢性病共存，需规范用药管理'
    ]
    for concern in concerns:
        doc.add_paragraph(f'⚠️ {concern}', style='List Number')
    
    # 六、后续健康建议
    add_heading(doc, '六、后续健康建议', level=1)
    suggestions = [
        '饮食管理 — 继续维持当前饮食方案，控制总热量和盐摄入（<5g/日）',
        '监测习惯 — 规律监测血糖血压，每周至少 3 次完整记录',
        '用药依从 — 按医嘱规范用药，不随意调整胰岛素剂量（诺和锐 6-6-4U）',
        '定期复查 — 每月复查肾功能、血常规、肝功能，关注血钾变化',
        '适量运动 — 每日散步 30 分钟，避免剧烈运动',
        '及时就医 — 出现不适及时就医，特别是低血糖、水肿复发等情况'
    ]
    for suggestion in suggestions:
        doc.add_paragraph(suggestion, style='List Number')
    
    # 七、每周数据汇总
    add_heading(doc, '七、每周数据汇总', level=1)
    weekly_headers = ['周次', '体重 (kg)', '空腹达标', '餐后达标', '关键事件']
    weekly_data = [
        ['第 1 周', '61.45', '7/7', '6/7', '初期水肿'],
        ['第 2 周', '60.0', '6/7', '6/7', '忘记胰岛素'],
        ['第 3 周', '58.5', '6/6', '3/7', '白内障手术'],
        ['第 4 周', '57.5', '6/6', '4/7', '频繁低血糖'],
        ['第 5 周', '57.5', '7/7', '7/7', '换德古胰岛素'],
        ['第 6 周', '58.1', '7/7', '5/7', '状态稳定'],
        ['第 7 周', '57.5', '6/7', '7/7', '高钾复查'],
        ['第 8 周', '56.45', '6/7', '5/7', '状态好转'],
        ['第 9 周', '56.1', '7/7', '3/7', '餐后波动']
    ]
    add_table(doc, weekly_data, weekly_headers)
    
    # 八、当前用药方案
    add_heading(doc, '八、当前用药方案', level=1)
    doc.add_paragraph('胰岛素：诺和锐 早 6U、午 6U、晚 4U')
    doc.add_paragraph('慢性病用药：降压、降脂、护肾、补血等 10+ 种')
    doc.add_paragraph('• 复方α-酮酸片（开同）', style='List Bullet')
    doc.add_paragraph('• 托伐普坦片', style='List Bullet')
    doc.add_paragraph('• 恩那度司他片', style='List Bullet')
    doc.add_paragraph('• 骨化三醇软胶囊', style='List Bullet')
    doc.add_paragraph('• 阿托伐他汀', style='List Bullet')
    doc.add_paragraph('• 阿司匹林', style='List Bullet')
    doc.add_paragraph('• 硫酸氢氯吡格雷', style='List Bullet')
    doc.add_paragraph('• 单硝酸异山梨酯', style='List Bullet')
    
    # 页脚信息
    doc.add_paragraph()
    doc.add_paragraph('_' * 50)
    footer = doc.add_paragraph()
    footer_run = footer.add_run('报告生成时间：2026-03-09\n')
    footer_run.italic = True
    footer_run2 = footer.add_run('温馨提示：本报告仅供参考，具体治疗方案请遵医嘱\n')
    footer_run2.italic = True
    footer_run3 = footer.add_run('下次复查建议：2026 年 4 月上旬（内分泌科、肾内科、心内科、眼科）')
    footer_run3.italic = True
    
    # 保存文档
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    import sys
    markdown_path = sys.argv[1] if len(sys.argv) > 1 else '/root/.openclaw/workspace/skills/health-weekly-report/output/王茂清 -9 周调理总结 + 体检整合报告.md'
    output_path = sys.argv[2] if len(sys.argv) > 2 else '/root/.openclaw/workspace/skills/health-weekly-report/output/王茂清 -9 周调理总结 + 体检整合报告.docx'
    
    result = generate_word_report(markdown_path, output_path)
    print(f"✅ Word 文档已生成：{result}")
